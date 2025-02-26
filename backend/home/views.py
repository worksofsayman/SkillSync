import os
from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect
from django.http import JsonResponse
import google.generativeai as genai
from docx import Document
from PyPDF2 import PdfReader
from django.conf import settings
from django.core.files.base import ContentFile
import asyncio
import aiohttp
from django.shortcuts import render
from django.views import View
from asgiref.sync import sync_to_async
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Resume
from .utils import convert_docx_to_pdf  # Assuming you already have this function for DOCX
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
import json
from fpdf import FPDF
import io
from .models import Resume
# Configure Generative AI API
genai.configure(api_key="AIzaSyAMHEgaEy6cAJK0bfziGAcLohnbECPyQE8")
model = genai.GenerativeModel("gemini-1.5-flash")

# Utility to read file content
def read_file_content(file_path, file_type):
    content = ""
    if file_type == "pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            content += page.extract_text()
    elif file_type in ["docx", "doc"]:
        doc = Document(file_path)
        for para in doc.paragraphs:
            content += para.text + "\n"
    elif file_type == "txt":
        with open(file_path, 'r') as file:
            content = file.read()
    else:
        raise ValueError("Unsupported file format.")
    return content

# Views
def index(request):
    return render(request, 'index.html')  # Serve the React build file

def suggest_resume(request):
    if request.method == 'POST' and request.FILES['resume']:
        # Save the uploaded resume
        resume_file = request.FILES['resume']
        fs = FileSystemStorage()
        filename = fs.save(resume_file.name, resume_file)
        file_path = fs.path(filename)

        # Determine file type
        _, file_extension = os.path.splitext(filename)
        file_type = file_extension.lower().strip('.')

        try:
            # Read file content
            resume_content = read_file_content(file_path, file_type)

            # Generate suggestions using Generative AI
            response = model.generate_content([
                "Suggest improvements for this resume:", resume_content
            ])
            suggestions = response.text
        except Exception as e:
            return render(request, 'suggest.html', {"error": f"Error processing file: {str(e)}"})

        return render(request, 'suggest.html', {"suggestions": suggestions})

    return render(request, 'suggest.html', {"error": "Please upload a valid file."})

def convert_to_cv(request):
    if request.method == 'POST' and request.FILES['resume']:
        # Save the uploaded resume
        resume_file = request.FILES['resume']
        fs = FileSystemStorage()
        filename = fs.save(resume_file.name, resume_file)
        file_path = fs.path(filename)

        # Determine file type
        _, file_extension = os.path.splitext(filename)
        file_type = file_extension.lower().strip('.')

        try:
            # Read file content
            resume_content = read_file_content(file_path, file_type)

            # Create a CV using python-docx
            doc = Document()
            doc.add_heading('Curriculum Vitae', level=1)
            doc.add_paragraph(resume_content)

            # Save the CV to the server
            cv_filename = f"CV_{os.path.splitext(filename)[0]}.docx"
            cv_path = os.path.join(settings.MEDIA_ROOT, cv_filename)
            doc.save(cv_path)

        except Exception as e:
            return render(request, 'cv_converter.html', {"error": f"Error processing file: {str(e)}"})

        return render(request, 'cv_converter.html', {"cv_url": f"/media/{cv_filename}"})

    return render(request, 'cv_converter.html', {"error": "Please upload a valid file."})



class USAJobsView(View):
    # Asynchronous method to fetch jobs from an API
    async def fetch_jobs(self, api_key, params):
        API_URL = 'https://data.usajobs.gov/api/search'

        headers = {
            'Host': 'data.usajobs.gov',
            'User-Agent': 'utkarshkushwaha246@gmail.com',
            'Authorization-Key': api_key
        }

        async with aiohttp.ClientSession() as session:
            try:
                async with session.get(API_URL, headers=headers, params=params, timeout=10) as response:
                    response.raise_for_status()  # Raise exception for errors
                    jobs = await response.json()
                    return jobs.get('SearchResult', {}).get('SearchResultItems', [])
            except Exception as e:
                return None

    # Asynchronous method to run both API requests in parallel
    async def get_jobs_data(self):
        params = {
            'Keyword': 'internship', 
            'LocationName': 'remote',  
            'ResultsPerPage': 10,  
        }

        api_key_1 = '3WGgHaa+SeqrF14ap76WXFHc1znhD8stg7bEJhkC6Y8='  # First API key
        api_key_2 = '2af21edf-fc6b-4d13-9cd8-17bba4a7ce1d'  # Second API key

        # Run both API calls concurrently
        job_data_1, job_data_2 = await asyncio.gather(
            self.fetch_jobs(api_key_1, params),
            self.fetch_jobs(api_key_2, params)
        )

        return job_data_1, job_data_2

    # Synchronous Django view method that will await async job fetching
    async def get(self, request):
        # Await the asynchronous task to fetch job data
        job_data_1, job_data_2 = await self.get_jobs_data()

        job_listings = []
        # Combine the results from both APIs if they are available
        if job_data_1:
            job_listings.extend([
                {
                    'title': job['MatchedObjectDescriptor']['PositionTitle'],
                    'organization': job['MatchedObjectDescriptor']['OrganizationName'],
                    'location': job['MatchedObjectDescriptor']['PositionLocationDisplay'],
                    'url': job['MatchedObjectDescriptor']['PositionURI'],
                }
                for job in job_data_1
            ])
        
        if job_data_2:
            job_listings.extend([
                {
                    'title': job['MatchedObjectDescriptor']['PositionTitle'],
                    'organization': job['MatchedObjectDescriptor']['OrganizationName'],
                    'location': job['MatchedObjectDescriptor']['PositionLocationDisplay'],
                    'url': job['MatchedObjectDescriptor']['PositionURI'],
                }
                for job in job_data_2
            ])
        
        return render(request, 'jobintern.html', {'jobs': job_listings})


from django.shortcuts import render
from .utils import Scrap_Internshala

def fetch_internships_view(request):
    if request.method == "POST":  
        skill = request.POST.get('skill', '')
        if skill:
            base_url = "https://internshala.com"
            try:
                internships_df = Scrap_Internshala(base_url, skill)
                internships = internships_df.to_dict(orient='records')  
                return render(request, 'internships.html', {'internships': internships, 'skill': skill})
            except Exception as e:
                return render(request, 'internships.html', {'error': str(e)})
        else:
            return render(request, 'internships.html', {'error': 'No skill provided'})
    return render(request, 'internships.html')  


from django.shortcuts import render
from .utils import scrap_jobs

def fetch_jobs_view(request):
    if request.method == "POST":  
        skill = request.POST.get('skill', '')
        if skill:
            base_url = "https://internshala.com"
            try:
                jobs_df = scrap_jobs(base_url, skill)
                jobs = jobs_df.to_dict(orient='records')  
                return render(request, 'jobs.html', {'jobs': jobs, 'skill': skill})
            except Exception as e:
                return render(request, 'jobs.html', {'error': str(e)})
        else:
            return render(request, 'jobs.html', {'error': 'No skill provided'})
    return render(request, 'jobs.html')





import openai
import requests
from django.shortcuts import render

# API Keys (Replace with your actual keys)
SERPAPI_KEY = "b1b74ec92d0858703b2a9bb9b5e0ac21746fd3bae2b9359dc6cedd581008242e"
OPENAI_API_KEY = "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"

def fetch_jobs_from_google(skill):
    """Fetch job listings from Google using SerpAPI"""
    url = "https://serpapi.com/search"
    params = {
        "engine": "google_jobs",
        "q": f"{skill} jobs",
        "hl": "en",
        "api_key": SERPAPI_KEY,
    }

    response = requests.get(url, params=params)
    data = response.json()
    
    jobs = []
    if "jobs_results" in data:
        for job in data["jobs_results"]:
            jobs.append({
                "title": job.get("title", "N/A"),
                "company": job.get("company_name", "N/A"),
                "location": job.get("location", "N/A"),
                "link": job.get("link", "#")
            })

    return jobs

def refine_jobs_with_chatgpt(jobs):
    """Use OpenAI API to refine and structure job data"""
    openai.api_key = OPENAI_API_KEY
    prompt = f"Format the following job listings into a structured table with site link:\n\n{jobs}"

    response = openai.ChatCompletion.create(
        model="gpt-4-turbo",
        messages=[{"role": "user", "content": prompt}]
    )

    return response["choices"][0]["message"]["content"]

def job_search_view(request):
    jobs = []
    skill = ""

    if request.method == "POST":
        skill = request.POST.get("skill", "")
        if skill:
            try:
                jobs = fetch_jobs_from_google(skill)
                formatted_jobs = refine_jobs_with_chatgpt(jobs)
            except Exception as e:
                return render(request, "ijobs.html", {"error": str(e)})

    return render(request, "ijobs.html", {"jobs": jobs, "skill": skill})

#KaustubhSen
def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password1 = request.POST['password1']
        password2 = request.POST['password2']

        if password1 != password2:
            messages.error(request, "Passwords do not match!")
            return redirect('register')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already taken!")
            return redirect('register')

        if User.objects.filter(email=email).exists():
            messages.error(request, "Email already registered!")
            return redirect('register')

        # Create the user
        user = User.objects.create_user(username=username, email=email, password=password1)
        user.save()
        messages.success(request, "Registration successful! Please log in.")
        return redirect('login')
    return render(request, 'register.html')


def login_user(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']

        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, "Login successful!")
            return redirect('home')  # Replace with your dashboard or home page
        else:
            messages.error(request, "Invalid username or password!")
            return redirect('home')
    return render(request, 'index.html')



def logout_user(request):
    logout(request)
    messages.success(request, "You have been logged out.")
    return redirect('login')


def upload_resume(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('resume')  # Get the uploaded file

        if uploaded_file:
            file_type = uploaded_file.name.split('.')[-1].lower()

            if file_type not in ['pdf', 'docx', 'jpeg', 'jpg', 'png', 'tiff']:
                messages.error(request, "Only pdf, docx, jpeg, jpg, png, and tiff files are allowed.")
                return redirect('upload')

            # Convert DOCX to PDF if necessary
            if file_type == 'docx':
                try:
                    pdf_content = convert_docx_to_pdf(uploaded_file)
                except Exception as e:
                    messages.error(request, f"Error processing DOCX file: {e}")
                    return redirect('upload')
            elif file_type in ['jpeg', 'jpg', 'png', 'tiff']:
                try:
                    resume_instance = Resume()  # Create an instance of Resume (you don't need to save yet)
                    pdf_content = resume_instance.convert_itp(uploaded_file)
                except Exception as e:
                    messages.error(request, f"Error processing image file: {e}")
                    return redirect('upload')
            else:
                pdf_content = uploaded_file.read()  # Read PDF directly

            # Save the PDF content in the database
            Resume.objects.create(user=request.user, resume=pdf_content) 
            messages.success(request, "Resume uploaded successfully!")
            return redirect('upload')
        else:
            messages.error(request, "No file uploaded.")
    return render(request, 'upload.html')


def download_resume(request, resume_id):
    # Fetch the resume instance from the database
    resume_instance = get_object_or_404(Resume, id=resume_id)
    
    if resume_instance.resume:
        # Create a response object to serve the file
        response = HttpResponse(resume_instance.resume, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="resume_{resume_id}.pdf"'
        return response
    else:
        return HttpResponse("No file found.", status=404)



def resume_editor(request):
    return render(request, 'resume_editor.html')


@login_required
def save_resume_pdf(request):
    if request.method == "POST":
        resume_text = "hello world "  # Replace with actual text from request.POST

        if not resume_text:
            return JsonResponse({"message": "Error: No resume text provided."}, status=400)

        try:
            # Create PDF using FPDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, resume_text)

            # Get binary data of PDF
            pdf_binary = pdf.output(dest='S').encode('latin1')  # Properly get binary content

            # Save to database
            filename = f"{request.user.username}_resume.pdf"
            resume = Resume.objects.create(user=request.user, resume=pdf_binary, filename=filename)

            return JsonResponse({"message": "Resume saved successfully!", "filename": filename})

        except Exception as e:
            return JsonResponse({"message": f"Error: {str(e)}"}, status=500)

    return render(request, 'save_resume_pdf.html')